#!/usr/bin/env python3
"""
Convert AWS_SETUP_GUIDE.md to Word document (.docx)
Preserves all content, formatting, headers, code blocks, and tables
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_code_block(doc, code_text):
    """Add formatted code block to document"""
    p = doc.add_paragraph(code_text, style='List Bullet')
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    # Add background color
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shading_elm)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(0)

def add_table_from_markdown(doc, table_text):
    """Parse and add markdown table to document"""
    lines = table_text.strip().split('\n')
    if len(lines) < 2:
        return

    # Parse header
    header_cells = [cell.strip() for cell in lines[0].split('|')[1:-1]]
    # Skip separator line
    rows_data = []
    for line in lines[2:]:
        if line.strip():
            row_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows_data.append(row_cells)

    # Create table
    table = doc.add_table(rows=1, cols=len(header_cells))
    table.style = 'Light Grid Accent 1'

    # Add header
    header_cells_obj = table.rows[0].cells
    for i, header in enumerate(header_cells):
        header_cells_obj[i].text = header
        for paragraph in header_cells_obj[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add rows
    for row_data in rows_data:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text

def convert_markdown_to_docx(md_file, output_file):
    """Convert markdown file to Word document"""
    doc = Document()

    # Read markdown file
    with open(md_file, 'r') as f:
        content = f.read()

    # Split into lines
    lines = content.split('\n')

    i = 0
    in_code_block = False
    code_block = []

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block = []
            else:
                in_code_block = False
                code_text = '\n'.join(code_block)
                add_code_block(doc, code_text)
                code_block = []
            i += 1
            continue

        if in_code_block:
            code_block.append(line)
            i += 1
            continue

        # Handle headers
        if line.startswith('# '):
            text = line.replace('# ', '').strip()
            heading = doc.add_heading(text, level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(12)
            i += 1
            continue

        if line.startswith('## '):
            text = line.replace('## ', '').strip()
            heading = doc.add_heading(text, level=2)
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        if line.startswith('### '):
            text = line.replace('### ', '').strip()
            heading = doc.add_heading(text, level=3)
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(8)
            i += 1
            continue

        if line.startswith('#### '):
            text = line.replace('#### ', '').strip()
            heading = doc.add_heading(text, level=4)
            heading.paragraph_format.space_before = Pt(6)
            heading.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Handle tables
        if line.strip().startswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            table_text = '\n'.join(table_lines)
            add_table_from_markdown(doc, table_text)
            continue

        # Handle bullet points and numbered lists
        if line.strip().startswith('- '):
            text = line.replace('- ', '').strip()
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        if line.strip().startswith('[ ]'):
            text = line.replace('[ ]', '☐').strip()
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        if line.strip().startswith('[x]') or line.strip().startswith('[✓]'):
            text = line.replace('[x]', '☑').replace('[✓]', '☑').strip()
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        # Handle line breaks and empty lines
        if not line.strip():
            i += 1
            continue

        # Handle bold text
        if line.strip().startswith('**'):
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.replace('**', ''))
                    run.bold = True
                elif part:
                    p.add_run(part)
            i += 1
            continue

        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph(line.strip())

        i += 1

    # Save document
    doc.save(output_file)
    print(f"✅ Converted to: {output_file}")

if __name__ == '__main__':
    convert_markdown_to_docx('AWS_SETUP_GUIDE.md', 'AWS_SETUP_GUIDE.docx')
